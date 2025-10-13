"""
Evidence Extractors

Content extraction classes for different file types commonly found
in legal evidence packages.
"""

import re
from pathlib import Path
from typing import Optional

import docx
from bs4 import BeautifulSoup


class BaseExtractor:
    """Base class for content extractors."""

    def extract(self, file_path: str) -> Optional[str]:
        """Extract text content from a file."""
        raise NotImplementedError("Subclasses must implement extract method")

    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""

        # Remove excessive whitespace
        text = re.sub(r"\s+", " ", text)

        # Remove control characters
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text)

        # Strip leading/trailing whitespace
        text = text.strip()

        return text


class DocumentExtractor(BaseExtractor):
    """Extractor for Microsoft Word documents (.docx)."""

    def extract(self, file_path: str) -> Optional[str]:
        """Extract text content from a DOCX file."""
        try:
            doc = docx.Document(file_path)

            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())

            # Combine all text
            full_text = "\n".join(paragraphs)

            return self.clean_text(full_text)

        except Exception as e:
            print(f"Error extracting from DOCX file {file_path}: {e}")
            return None


class HTMLExtractor(BaseExtractor):
    """Extractor for HTML files, including email HTML."""

    def extract(self, file_path: str) -> Optional[str]:
        """Extract text content from an HTML file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                html_content = f.read()

            # Parse HTML with BeautifulSoup
            soup = BeautifulSoup(html_content, "html.parser")

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Extract text content
            text = soup.get_text()

            return self.clean_text(text)

        except Exception as e:
            print(f"Error extracting from HTML file {file_path}: {e}")
            return None


class EmailExtractor(BaseExtractor):
    """Extractor specifically for email files (HTML format)."""

    def extract(self, file_path: str) -> Optional[str]:
        """Extract structured content from an email HTML file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                html_content = f.read()

            # Parse HTML with BeautifulSoup
            soup = BeautifulSoup(html_content, "html.parser")

            # Extract email metadata
            email_parts = []

            # Look for email headers (From, To, Subject, Date)
            headers = self._extract_email_headers(soup)
            if headers:
                email_parts.append("EMAIL HEADERS:")
                email_parts.extend(headers)
                email_parts.append("")

            # Extract main email body
            body = self._extract_email_body(soup)
            if body:
                email_parts.append("EMAIL BODY:")
                email_parts.append(body)

            # Combine all parts
            full_text = "\n".join(email_parts)

            return self.clean_text(full_text)

        except Exception as e:
            print(f"Error extracting from email file {file_path}: {e}")
            return None

    def _extract_email_headers(self, soup: BeautifulSoup) -> list:
        """Extract email headers from HTML."""
        headers = []

        # Common patterns for email headers
        header_patterns = [
            ("From:", r"From:?\s*(.+?)(?:<|$)"),
            ("To:", r"To:?\s*(.+?)(?:<|$)"),
            ("Subject:", r"Subject:?\s*(.+?)$"),
            ("Date:", r"(?:Sent|Date):?\s*(.+?)$"),
            ("CC:", r"Cc:?\s*(.+?)(?:<|$)"),
        ]

        text = soup.get_text()
        lines = text.split("\n")

        for line in lines:
            line = line.strip()
            for header_name, pattern in header_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    headers.append(f"{header_name} {match.group(1).strip()}")

        return headers

    def _extract_email_body(self, soup: BeautifulSoup) -> str:
        """Extract the main email body content."""
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Look for main content divs
        content_divs = soup.find_all(
            "div", class_=lambda x: x and "elementToProof" in x
        )

        if content_divs:
            # Extract text from content divs
            body_parts = []
            for div in content_divs:
                text = div.get_text().strip()
                if text and len(text) > 10:  # Filter out very short content
                    body_parts.append(text)
            return "\n".join(body_parts)
        else:
            # Fallback: extract all text
            return soup.get_text()


class PDFExtractor(BaseExtractor):
    """Extractor for PDF files (placeholder for future implementation)."""

    def extract(self, file_path: str) -> Optional[str]:
        """Extract text content from a PDF file."""
        # This would require additional dependencies like PyPDF2 or pdfplumber
        # For now, return a placeholder
        print(f"PDF extraction not yet implemented for {file_path}")
        return None


class TextExtractor(BaseExtractor):
    """Extractor for plain text files."""

    def extract(self, file_path: str) -> Optional[str]:
        """Extract content from a plain text file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            return self.clean_text(content)

        except Exception as e:
            print(f"Error extracting from text file {file_path}: {e}")
            return None


# Factory function for getting appropriate extractor
def get_extractor(file_path: str) -> Optional[BaseExtractor]:
    """Get the appropriate extractor for a file based on its extension."""
    file_path = Path(file_path)
    extension = file_path.suffix.lower().lstrip(".")

    extractors = {
        "docx": DocumentExtractor(),
        "html": HTMLExtractor(),
        "htm": HTMLExtractor(),
        "pdf": PDFExtractor(),
        "txt": TextExtractor(),
        "md": TextExtractor(),
    }

    return extractors.get(extension)
